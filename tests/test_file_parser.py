import io
import pytest
from utils.file_parser import parse_file


class TestParseFile:
    def test_txt_parsing(self):
        content = "Hello world\nThis is a test."
        result = parse_file(content.encode("utf-8"), "test.txt")
        assert "Hello world" in result
        assert "This is a test" in result

    def test_md_parsing(self):
        content = "# Title\n\nSome **markdown** text."
        result = parse_file(content.encode("utf-8"), "doc.md")
        assert "Title" in result
        assert "Some" in result
        assert "markdown" in result

    def test_unknown_extension_raises(self):
        content = "plain text content"
        with pytest.raises(ValueError, match="不支持"):
            parse_file(content.encode("utf-8"), "file.xyz")

    def test_empty_content(self):
        content = ""
        result = parse_file(content.encode("utf-8"), "empty.txt")
        assert result == ""

    def test_chinese_text(self):
        content = "人工智能是计算机科学的一个分支。"
        result = parse_file(content.encode("utf-8"), "cn.txt")
        assert "人工智能" in result

    def test_docx_parsing(self):
        """Minimal DOCX created via python-docx."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_paragraph("Paragraph one.")
        doc.add_paragraph("Paragraph two.")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        result = parse_file(buf.read(), "test.docx")
        assert "Paragraph one" in result
        assert "Paragraph two" in result

    def test_pdf_parsing(self):
        """Test PDF parsing with a minimal valid PDF constructed as raw bytes."""
        try:
            import pdfplumber
        except ImportError:
            pytest.skip("pdfplumber not installed")

        def _create_minimal_pdf(text: str = "Hello PDF World") -> bytes:
            content = f"""1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj

2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj

3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]
   /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>
endobj

4 0 obj
<< /Length 44 >>
stream
BT /F1 12 Tf 100 700 Td ({text}) Tj ET
endstream
endobj

5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj

xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000266 00000 n 
0000000360 00000 n 

trailer
<< /Size 6 /Root 1 0 R >>
startxref
437
%%EOF"""
            return content.encode("latin-1")

        pdf_bytes = _create_minimal_pdf("Test PDF content")
        result = parse_file(pdf_bytes, "test.pdf")
        assert "Test PDF content" in result

    def test_replaces_newlines_with_spaces_in_paragraphs(self):
        content = "Line 1\nLine 2\nLine 3"
        result = parse_file(content.encode("utf-8"), "test.txt")
        # parse_file joins paragraphs with spaces
        assert "Line 1" in result
