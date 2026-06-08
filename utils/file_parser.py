"""
文件解析模块 - 支持 PDF / Word / 纯文本 / Markdown 文件的文本提取
"""

import io
import re
from pathlib import Path

import pdfplumber
from docx import Document


def parse_file(file_bytes: bytes, filename: str) -> str:
    """根据文件扩展名自动选择解析方式，返回提取的纯文本内容。

    Args:
        file_bytes: 文件二进制内容
        filename: 原始文件名（用于判断扩展名）

    Returns:
        提取出的纯文本字符串

    Raises:
        ValueError: 不支持的文件类型
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _parse_pdf(file_bytes)
    elif ext == ".docx":
        return _parse_docx(file_bytes)
    elif ext == ".doc":
        raise ValueError(
            "不支持旧版 Word 格式 (.doc)。请将文件另存为 .docx 格式后再上传。\n"
            "操作方法：用 Word 打开文件 → 文件 → 另存为 → 选择「Word 文档(*.docx)」"
        )
    elif ext in (".txt", ".md", ".markdown"):
        return _parse_text(file_bytes)
    else:
        raise ValueError(f"不支持的文件类型: {ext}，目前支持 PDF、Word(.docx)、纯文本(.txt)、Markdown(.md)")


def _parse_pdf(file_bytes: bytes) -> str:
    """解析 PDF 文件"""
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return _clean_text("\n".join(text_parts))


def _parse_docx(file_bytes: bytes) -> str:
    """解析 Word 文档"""
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # 也提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(" | ".join(row_text))

    return _clean_text("\n".join(text_parts))


def _parse_text(file_bytes: bytes) -> str:
    """解析纯文本 / Markdown 文件（自动检测编码）"""
    # 尝试常见编码
    encodings = ["utf-8", "gbk", "gb2312", "utf-16", "latin-1"]
    for enc in encodings:
        try:
            text = file_bytes.decode(enc)
            return _clean_text(text)
        except (UnicodeDecodeError, UnicodeError):
            continue
    # 兜底：忽略无法解码的字符
    return _clean_text(file_bytes.decode("utf-8", errors="ignore"))


def _clean_text(text: str) -> str:
    """清洗文本：合并多余空白、去除控制字符"""
    # 去除空行过多的情况
    text = re.sub(r"\n{3,}", "\n\n", text)
    # 去除控制字符（保留换行和制表符）
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
    return text.strip()
